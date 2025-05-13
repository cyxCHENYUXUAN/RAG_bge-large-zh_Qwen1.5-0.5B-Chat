import os
import json
import re
from pathlib import Path
from typing import List, Dict, Any
import PyPDF2
from docx import Document
from datetime import datetime
import textract

def read_document(file_path: str) -> str:
    """读取不同类型的文档并提取文本内容"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == '.pdf':
        return read_pdf(file_path)
    elif suffix == '.docx':
        return read_docx(file_path)
    elif suffix == '.doc':
        return read_doc(file_path)
    elif suffix == '.txt':
        return read_txt(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {suffix}")

def read_pdf(file_path: str) -> str:
    """从PDF文件中提取文本"""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text() or ""
                # 添加页码信息
                text += f"\n--- 第{page_num + 1}页 ---\n{page_text}\n"
    except Exception as e:
        print(f"DOC_PROC: 读取PDF {file_path} 时出错: {e}")
        raise # Re-raise error
    return text

def read_docx(file_path: str) -> str:
    """从DOCX文件中提取文本"""
    # 多种方法依次尝试读取DOCX文件，确保最大兼容性
    all_methods_tried = False
    result_text = ""
    
    print(f"DOC_PROC: 开始读取DOCX文件: {file_path}")
    
    # 方法1: 使用python-docx库 (标准方法)
    try:
        doc = Document(file_path)
        paragraphs = []
        
        # 提取正文内容
        try:
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
        except Exception as para_error:
            print(f"DOC_PROC: 读取段落时出错: {para_error}，尝试继续处理")
        
        # 提取表格内容
        try:
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
        except Exception as table_error:
            print(f"DOC_PROC: 读取表格时出错: {table_error}，尝试继续处理")
        
        if paragraphs:
            result_text = "\n".join(paragraphs)
            print(f"DOC_PROC: DOCX文件读取成功(方法1)，共提取 {len(paragraphs)} 个段落")
            return result_text
        else:
            print("DOC_PROC: 方法1未提取到任何内容，尝试备用方法")
    except Exception as e:
        print(f"DOC_PROC: 主要读取方法出错: {e}，尝试备用方法")
    
    # 方法2: 使用zipfile直接解析XML (备用方法1)
    try:
        import zipfile
        from xml.etree.ElementTree import XML, ParseError
        
        print(f"DOC_PROC: 尝试使用备用方法1读取DOCX文件: {file_path}")
        WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        PARA = WORD_NAMESPACE + 'p'
        TEXT = WORD_NAMESPACE + 't'
        
        with zipfile.ZipFile(file_path) as zip_file:
            try:
                try:
                    xml_content = zip_file.read('word/document.xml')
                    tree = XML(xml_content)
                    
                    texts = []
                    for paragraph in tree.iter(PARA):
                        paragraph_text = ''.join(node.text for node in paragraph.iter(TEXT) if node.text)
                        if paragraph_text.strip():
                            texts.append(paragraph_text)
                    
                    if texts:
                        result_text = '\n'.join(texts)
                        print(f"DOC_PROC: 使用备用方法1成功读取DOCX，提取了 {len(texts)} 个段落")
                        return result_text
                    else:
                        print("DOC_PROC: 备用方法1未提取到任何文本")
                except KeyError:
                    print("DOC_PROC: 找不到word/document.xml文件")
                except ParseError as xml_error:
                    print(f"DOC_PROC: XML解析错误: {xml_error}")
            except Exception as backup_error:
                print(f"DOC_PROC: 备用方法1失败: {backup_error}")
    except ImportError:
        print("DOC_PROC: 无法导入备用方法1所需的模块")
    
    # 方法3：最后的尝试，简单地读取docx中的所有文本文件
    try:
        import zipfile
        
        print(f"DOC_PROC: 尝试使用备用方法2读取DOCX文件: {file_path}")
        texts = []
        
        with zipfile.ZipFile(file_path) as zip_file:
            try:
                for file_info in zip_file.infolist():
                    if file_info.filename.startswith('word/') and file_info.filename.endswith('.xml'):
                        try:
                            content = zip_file.read(file_info.filename).decode('utf-8', errors='ignore')
                            # 简单文本提取 - 移除所有XML标签
                            import re
                            text_only = re.sub(r'<[^>]+>', ' ', content)
                            text_only = re.sub(r'\s+', ' ', text_only).strip()
                            if text_only:
                                texts.append(text_only)
                        except Exception as file_error:
                            print(f"DOC_PROC: 读取 {file_info.filename} 时出错: {file_error}")
                
                if texts:
                    result_text = '\n'.join(texts)
                    print(f"DOC_PROC: 使用备用方法2成功读取DOCX，提取了 {len(texts)} 个文本块")
                    return result_text
                else:
                    print("DOC_PROC: 备用方法2未提取到任何文本")
            except Exception as e:
                print(f"DOC_PROC: 备用方法2遍历文件失败: {e}")
    except ImportError:
        print("DOC_PROC: 无法导入备用方法2所需的模块")
    
    # 所有方法都失败，返回一个空字符串，上层处理将创建一个默认文本块
    all_methods_tried = True
    if all_methods_tried and not result_text:
        print(f"DOC_PROC: 所有方法都无法从文件中提取文本: {file_path}")
        return ""

def read_doc(file_path: str) -> str:
    """从DOC文件中提取文本（旧版Word格式）"""
    print(f"DOC_PROC: 开始读取DOC文件: {file_path}")
    result_text = ""
    
    # 方法1: 使用textract库提取内容
    try:
        print(f"DOC_PROC: 尝试使用textract读取DOC文件")
        text = textract.process(file_path).decode('utf-8', errors='ignore')
        
        if text and len(text.strip()) > 0:
            result_text = text
            print(f"DOC_PROC: 使用textract成功读取DOC文件，提取了约 {len(text)} 个字符")
            return result_text
        else:
            print("DOC_PROC: textract提取的文本为空，尝试备用方法")
    except Exception as e:
        print(f"DOC_PROC: 使用textract读取出错: {e}，尝试备用方法")
    
    # 方法2: 尝试使用antiword (如果系统中安装了)
    try:
        import subprocess
        print(f"DOC_PROC: 尝试使用antiword读取DOC文件")
        
        result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
        if result.returncode == 0 and result.stdout:
            result_text = result.stdout
            print(f"DOC_PROC: 使用antiword成功读取DOC文件，提取了约 {len(result_text)} 个字符")
            return result_text
        else:
            print(f"DOC_PROC: antiword未能成功提取文本: {result.stderr}")
    except Exception as e:
        print(f"DOC_PROC: 使用antiword读取失败: {e}")
    
    # 方法3: 在Windows环境下尝试使用win32com调用MS Word
    if os.name == 'nt':  # 检查是否是Windows系统
        try:
            import win32com.client
            import pythoncom
            
            print(f"DOC_PROC: 尝试使用win32com/MS Word读取DOC文件")
            pythoncom.CoInitialize()
            
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            try:
                doc = word_app.Documents.Open(os.path.abspath(file_path))
                result_text = doc.Content.Text
                doc.Close()
            finally:
                word_app.Quit()
                
            if result_text:
                print(f"DOC_PROC: 使用win32com成功读取DOC文件，提取了约 {len(result_text)} 个字符")
                return result_text
            else:
                print("DOC_PROC: win32com提取的文本为空")
        except Exception as e:
            print(f"DOC_PROC: 使用win32com读取失败: {e}")
    
    # 如果所有方法都失败
    if not result_text:
        print(f"DOC_PROC: 所有方法都无法从DOC文件中提取文本: {file_path}")
        return ""

def read_txt(file_path: str) -> str:
    """从TXT文件中提取文本"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"DOC_PROC: 读取TXT {file_path} 时出错: {e}")
        raise

def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
    """将文本分割成有重叠的块，智能识别段落和语义边界"""
    chunks = []
    if not text:
        return chunks
        
    # 按段落分割文本 (保守方式，兼容多种换行)
    paragraphs = re.split(r'\n[\n\s]*\n', text) # Split on blank lines (possibly with spaces)
    
    current_chunk = ""
    current_size = 0
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
            
        paragraph_size = len(paragraph)
        
        # 如果段落本身就很大
        if paragraph_size > chunk_size:
            # Further split large paragraphs by sentences (or simply force split if needed)
            sentences = re.split(r'(?<=[.!?])\s+|(?<=[。！？])\s+', paragraph) # Add Chinese punctuation
            temp_paragraph_chunk = ""
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                sentence_size = len(sentence)

                # If a single sentence is too large, add it as its own chunk
                if sentence_size > chunk_size:
                    if temp_paragraph_chunk:
                        chunks.append({"text": temp_paragraph_chunk.strip()})
                        temp_paragraph_chunk = get_overlap_text(temp_paragraph_chunk, overlap) # Start new with overlap
                    chunks.append({"text": sentence}) # Add the huge sentence
                    continue # Move to the next sentence
                
                # If adding the sentence exceeds chunk size, finalize previous and start new
                if len(temp_paragraph_chunk) + sentence_size > chunk_size and temp_paragraph_chunk:
                    chunks.append({"text": temp_paragraph_chunk.strip()})
                    temp_paragraph_chunk = get_overlap_text(temp_paragraph_chunk, overlap) + sentence + " "
                else:
                    temp_paragraph_chunk += sentence + " "
            
            # Add any remaining part of the large paragraph
            if temp_paragraph_chunk.strip():
                 # Check if adding this last part to the *overall* current_chunk exceeds limit
                 if current_size + len(temp_paragraph_chunk) > chunk_size and current_chunk:
                    chunks.append({"text": current_chunk.strip()})
                    current_chunk = get_overlap_text(current_chunk, overlap)
                    current_size = len(current_chunk)
                 current_chunk += temp_paragraph_chunk.strip() + "\n\n"
                 current_size += len(temp_paragraph_chunk) + 2

        # If paragraph is not too large itself
        else:
            # 检查添加段落后是否超过块大小
            if current_size + paragraph_size > chunk_size and current_chunk:
                chunks.append({"text": current_chunk.strip()})
                # 保留一部分重叠内容
                overlap_text = get_overlap_text(current_chunk, overlap)
                current_chunk = overlap_text + paragraph + "\n\n"
                current_size = len(overlap_text) + paragraph_size + 2
            else:
                current_chunk += paragraph + "\n\n"
                current_size += paragraph_size + 2
    
    # 添加最后一块（如果有内容）
    if current_chunk.strip():
        chunks.append({"text": current_chunk.strip()})
    
    # Add chunk index metadata right after splitting
    for i, chunk_dict in enumerate(chunks):
        chunk_dict["chunk_index"] = i
        
    print(f"DOC_PROC: 将文本分割成 {len(chunks)} 个块.")
    return chunks

def get_overlap_text(text: str, overlap_size: int) -> str:
    """从文本末尾获取指定大小的重叠内容 (基于字符)"""
    if not text or overlap_size <= 0:
        return ""
    # Get last N characters as overlap
    overlap_text = text[-overlap_size:]
    # Try to find a space to make the overlap cleaner (optional)
    last_space = overlap_text.rfind(' ')
    if last_space > overlap_size // 2: # Only cut if it leaves reasonable overlap
        overlap_text = overlap_text[last_space+1:]
        
    return overlap_text + " " # Add space for separation

def process_document(original_file_path: str, processed_output_dir: Path) -> str:
    """处理文档（读取、分块）并将结果（文本块）保存为JSON文件到指定目录"""
    try:
        print(f"DOC_PROC: 开始处理文档: {original_file_path}")
        text = read_document(original_file_path)
        
        # 检查是否获取到有效文本
        if not text or len(text.strip()) == 0:
            print(f"DOC_PROC: 从文件 {original_file_path} 中提取的文本为空")
            # 创建一个包含警告信息的默认块
            chunks = [{"text": "此文档未能提取到有效文本内容。可能是文档为空、格式不兼容或仅包含无法读取的内容（如图片）。", "chunk_index": 0}]
        else:
            chunks = split_text_into_chunks(text)
            if not chunks:
                print(f"DOC_PROC: 文本分块结果为空，创建默认块")
                chunks = [{"text": "此文档的文本已被提取，但无法有效分块。可能是文档格式特殊或内容过少。", "chunk_index": 0}]
        
        original_path_obj = Path(original_file_path)
        document_info = {
            "original_filename": original_path_obj.name, # Store original filename
            "processed_at": datetime.now().isoformat(),
            "chunks": chunks # List of dicts, each dict is {'text': ..., 'chunk_index': ...}
        }
        
        # 确保输出目录存在
        processed_output_dir.mkdir(parents=True, exist_ok=True)
        
        # 使用原始文件名的stem创建输出文件名，并确保其安全性
        stem = original_path_obj.stem
        # 移除可能导致JSON路径问题的字符
        safe_stem = ''.join(c if c.isalnum() or c in '-_' else '_' for c in stem)
        output_filename = f"{safe_stem}_processed.json"
        output_path = processed_output_dir / output_filename
        
        print(f"DOC_PROC: 处理完成，将结果保存到: {output_path}")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(document_info, f, ensure_ascii=False, indent=2)
        
        return str(output_path)
    except ValueError as ve:
        print(f"DOC_PROC: 处理文档 {original_file_path} 时出错 (不支持的文件类型?): {ve}")
        raise
    except Exception as e:
        print(f"DOC_PROC: 处理文档 {original_file_path} 时发生意外错误: {e}")
        import traceback
        traceback.print_exc()
        raise

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 2:
        input_file = sys.argv[1]
        output_dir = Path(sys.argv[2])
        print(f"DOC_PROC: 通过命令行测试处理: {input_file} -> {output_dir}")
        try:
            output_json_path = process_document(input_file, output_dir)
            print(f"DOC_PROC: 文档处理完成，JSON保存到: {output_json_path}")
        except Exception as e:
            print(f"DOC_PROC: 命令行测试处理失败: {e}")
    else:
        print("DOC_PROC CLI Usage: python document_processor.py <input_file_path> <output_directory>") 